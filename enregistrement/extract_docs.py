from docx import Document

# Lire SDAD.docx (thèmes proposés par profs)
with open('g:/AZ/Documents/gestion des appelles telephoniques/enregistrement/extracted_sdad.txt', 'w', encoding='utf-8') as f:
    f.write('='*60 + '\n')
    f.write('FICHIER 1: SDAD.docx (Thèmes proposés par les profs)\n')
    f.write('='*60 + '\n')
    doc1 = Document(r'g:\AZ\Documents\gestion des appelles telephoniques\enregistrement\SDAD.docx')
    for p in doc1.paragraphs:
        if p.text.strip():
            f.write(p.text + '\n')

# Lire FINAL_RV (28).docx (exemple thèse encadrant)  
with open('g:/AZ/Documents/gestion des appelles telephoniques/enregistrement/extracted_thesis.txt', 'w', encoding='utf-8') as f:
    f.write('='*60 + '\n')
    f.write('FICHIER 2: FINAL_RV (28).docx (Exemple thèse encadrant)\n')
    f.write('='*60 + '\n')
    doc2 = Document(r'g:\AZ\Documents\gestion des appelles telephoniques\enregistrement\FINAL_RV (28).docx')
    for i, p in enumerate(doc2.paragraphs[:150]):
        if p.text.strip():
            f.write(p.text + '\n')

print("Fichiers extraits avec succes!")
